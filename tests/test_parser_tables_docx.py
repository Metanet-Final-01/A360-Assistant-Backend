"""파서 강화 테스트 (RPA-42) — PDF 구조화 표, DOCX 문단·표, 자연어 텍스트."""

import io

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

from app.services.parser import parse_document, parse_text
from app.services.parser.pdf import _clean_table


# --- DOCX ---

def _make_docx() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("금 시세 조회 자동화")
    doc.add_paragraph("네이버 증권에서 국내 금 시세를 조회한다")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "단계"
    table.cell(0, 1).text = "시스템"
    table.cell(1, 0).text = "시세 조회"
    table.cell(1, 1).text = "Edge"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_extracts_text_and_table_in_order():
    result = parse_document("업무.docx", _make_docx())

    assert result["parser"] == "python-docx"
    assert result["page_count"] == 1
    blocks = result["pages"][0]["blocks"]
    types = [b["type"] for b in blocks]
    assert types == ["text", "text", "table"]  # 문단 2개 뒤에 표 (등장 순서 보존)

    assert blocks[0]["text"] == "금 시세 조회 자동화"
    table = blocks[2]
    assert table["rows"][0] == ["단계", "시스템"]
    assert table["rows"][1] == ["시세 조회", "Edge"]

    assert "네이버 증권" in result["full_text"]
    assert "시세 조회 | Edge" in result["full_text"]


def test_docx_merged_cells_not_duplicated():
    """가로 병합 셀은 row.cells가 같은 셀을 반복 반환 — _tc로 중복 제거되어야 한다."""
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=3)
    table.cell(1, 0).text = "값1"
    table.cell(1, 1).text = "값2"
    table.cell(1, 2).text = "값3"
    # 첫 행 3칸을 하나로 병합 → row.cells는 같은 셀을 3번 반환
    merged = table.cell(0, 0).merge(table.cell(0, 2))
    merged.text = "제목"

    buf = io.BytesIO()
    doc.save(buf)
    result = parse_document("merged.docx", buf.getvalue())

    table_block = next(b for b in result["pages"][0]["blocks"] if b["type"] == "table")
    assert table_block["rows"][0] == ["제목"]  # 3번 아니라 1번
    assert table_block["rows"][1] == ["값1", "값2", "값3"]


def test_docx_empty_warns():
    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_document("빈.docx", buf.getvalue())
    assert result["pages"][0]["blocks"] == []
    assert result["warnings"]


# --- PDF 구조화 표 (pdfplumber) ---

def _make_pdf_with_table() -> bytes:
    # reportlab 기본 폰트(Helvetica)는 한글 글리프가 없어 PDF 픽스처는 영문으로 만든다.
    # 한글 표 추출은 DOCX/PPTX 테스트가 커버하고, 여기선 PDF의 '구조화 표 검출'만 검증한다.
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter)
    data = [["Step", "System"], ["Search", "Edge"], ["Write", "Excel"]]
    table = Table(data, colWidths=[2 * inch, 2 * inch], rowHeights=0.4 * inch)
    # 테두리(선) — pdfplumber는 기본적으로 선 기반으로 표를 검출한다
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    pdf.build([table])
    return buf.getvalue()


def test_pdf_extracts_structured_table():
    result = parse_document("table.pdf", _make_pdf_with_table())

    assert "tables" in result["parser"]  # "pypdf+tables"
    table_blocks = [
        b for page in result["pages"] for b in page["blocks"] if b["type"] == "table"
    ]
    assert table_blocks, "구조화된 표 블록이 있어야 한다"
    rows = table_blocks[0]["rows"]
    assert ["Step", "System"] in rows
    assert ["Search", "Edge"] in rows
    # 표 내용은 full_text에 있다 (구조화 표 → pipe 렌더)
    assert "Search" in result["full_text"]
    # 중복 방지: 표 셀 값이 텍스트 블록에도 또 들어가면 안 된다 (analyze는 blocks를 읽음)
    text_blocks = [
        b for page in result["pages"] for b in page["blocks"] if b["type"] == "text"
    ]
    text_blob = " ".join(b["text"] for b in text_blocks)
    assert "Search" not in text_blob, "표 셀이 텍스트 블록에 중복되면 안 됨"


def test_clean_table_normalizes_none_and_blank_rows():
    raw = [["a", None, "b"], [None, None, None], ["", "c", ""]]
    assert _clean_table(raw) == [["a", "", "b"], ["", "c", ""]]


# --- 자연어 텍스트 입력 ---

def test_parse_text_wraps_natural_language():
    result = parse_text("  웹에서 금 시세를 긁어 엑셀로 정리해줘  ")
    assert result["parser"] == "text"
    assert result["page_count"] == 1
    assert result["full_text"] == "웹에서 금 시세를 긁어 엑셀로 정리해줘"
    assert result["pages"][0]["blocks"][0]["text"].startswith("웹에서 금 시세")


def test_parse_text_empty_warns():
    result = parse_text("   ")
    assert result["pages"][0]["blocks"] == []
    assert result["warnings"]
