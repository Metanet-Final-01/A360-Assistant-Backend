"""DOCX 텍스트·표 추출 — 문단과 표를 본문 등장 순서 그대로 읽는다.

python-docx의 .paragraphs/.tables는 순서 정보가 없어 본문 XML을 직접 순회한다
(문단 사이에 낀 표의 읽기 순서를 보존해야 분석 품질이 유지됨). DOCX는 페이지 개념이
런타임에만 있으므로 단일 페이지(page=1)로 취급한다.
"""

import io
import logging

from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


def _iter_block_items(parent):
    """본문 최상위 블록(문단/표)을 등장 순서대로 순회한다."""
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_docx(content: bytes) -> dict:
    document = DocxDocument(io.BytesIO(content))

    blocks: list[dict] = []
    full_parts: list[str] = []
    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text:
                blocks.append({"type": "text", "text": text})
                full_parts.append(text)
        else:  # Table
            rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            rows = [r for r in rows if any(r)]  # 완전 빈 행 제거
            if rows:
                blocks.append({"type": "table", "rows": rows})
                full_parts.append("\n".join(" | ".join(r) for r in rows))

    warnings: list[str] = []
    if not blocks:
        warnings.append("문서에서 텍스트를 찾지 못함 (이미지 전용/빈 문서 가능성)")

    return {
        "parser": "python-docx",
        "page_count": 1,
        "pages": [{"page": 1, "blocks": blocks}],
        "full_text": "\n\n".join(full_parts),
        "warnings": warnings,
    }
